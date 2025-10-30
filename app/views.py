from django.shortcuts import render, redirect
from .models import Rubric, Student, Progress
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_http_methods
from django.http import JsonResponse
from django.db.models import Avg
from django.contrib import messages
from datetime import datetime, timedelta
import json
import google.generativeai as genai
from django.conf import settings
from docx import Document
from PyPDF2 import PdfReader
import io

# Configure Gemini client once
genai.configure(api_key=settings.GOOGLE_API_KEY)

def grade_assignment(request):
    model = genai.GenerativeModel("gemini-2.0-flash-lite")
    student = request.POST.get('student', '').strip()
    rubric_id = request.POST.get('rubric', '').strip()
    rubric = Rubric.objects.get(id=rubric_id)
    strictness = rubric.strictness
    grade_level = rubric.grade_level
    training_data = rubric.training_data if rubric.training_data else ""

    # Handle file upload if present
    if 'assignment_file' in request.FILES:
        file = request.FILES['assignment_file']
        
        # Check file type
        if file.name.endswith('.docx'):
            # Handle Word document
            doc = Document(io.BytesIO(file.read()))
            description = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        elif file.name.endswith('.pdf'):
            # Handle PDF files
            pdf = PdfReader(io.BytesIO(file.read()))
            text_content = []
            for page in pdf.pages:
                text_content.append(page.extract_text())
            description = '\n'.join(text_content)
        elif file.name.endswith('.txt'):
            # Handle text files
            description = file.read().decode('utf-8')
        else:
            return JsonResponse({
                'status': 'error',
                'message': 'Unsupported file type. Please upload .txt, .docx, or .pdf files only.'
            }, status=400)
    else:
        description = request.POST.get('assignment-description', '').strip()

    # Format training data as examples if it exists
    training_examples_str = ""
    if training_data:
        examples = []
        for i, example in enumerate(training_data[:3], 1):  # Show max 3 examples
            examples.append(f"Example {i}:\nAssignment: {example['assignment'][:500]}...\nScore Given: {example['score']}/10")
        training_examples_str = "\n\nHere are some past graded examples to guide your grading style:\n" + "\n\n".join(examples)
    print(training_examples_str)
    prompt_score = f""" 
    Strictness rating (How strict you should grade this student's assignment): {strictness}
    Grade Level (The educational level of the student): {grade_level}
    Rubric (The criteria of which you will be grading the assignment): {rubric.description}
    
    {training_examples_str}
    
    Student's Assignment (The work that you will be grading): {description}
    
    You are a teacher grading an assignment, with a strictness rating of {strictness} / 10. Grade this students assignment with the 
    following rubric and only output with this format. Give the total score as an addition of all the scores of the different criteria. i want the only thing you tell me to be the score on this JSON format: {{
            "total_score": number,
            "percentage": number,
            "feedback_summary": string,
            "rubric_breakdown": [
                {{
                    "criterion": string,
                    "score": number,
                    "feedback": string
                }}
            ]
        }}"""
    assignment_score = model.generate_content(prompt_score).text
    print(assignment_score)
    # Clean the response - remove markdown code blocks
    if assignment_score.startswith('```json'):
        assignment_score = assignment_score.replace('```json', '').replace('```', '').strip()
    elif assignment_score.startswith('```'):
        assignment_score = assignment_score.replace('```', '').strip()
    
    try:
        score_data = json.loads(assignment_score)
        
        total_score = score_data.get('total_score')
        feedback_summary = score_data.get('feedback_summary')
        rubric_breakdown = score_data.get('rubric_breakdown', [])
        percentage = score_data.get('percentage')
        print(assignment_score)
        # Get student and rubric objects for context
        student_obj = Student.objects.get(id=student)
        # Prepare context for template
        context = {
            'student': student_obj,
            'rubric': rubric,
            'assignment_description': description,
            'total_score': total_score,
            'feedback_summary': feedback_summary,
            'rubric_breakdown': rubric_breakdown,
            'strictness': strictness,
            'grade_level': grade_level
        }
        prompt_evaluate = f"""
        Evaluate if this AI grading aligns with the rubric criteria:
        
        RUBRIC CRITERIA:
        {rubric.description}
        
        GRADING PARAMETERS:
        - Strictness: {strictness}/10
        - Grade Level: {grade_level}
        
        AI'S GRADING:
        - Total Score: {total_score}
        - Overall Feedback: {feedback_summary}
        - Criterion Scores: {[f"{c['criterion']}: {c['score']}" for c in rubric_breakdown]}
        
        PAST GRADING EXAMPLES (for reference):
        {training_examples_str if training_examples_str else 'No examples available'}
        
        TASK: Determine if the AI correctly applied the rubric criteria to grade this assignment.
        
        Respond ONLY in this exact format:
        PASS or FAIL
        
        Reasoning: [Your explanation]
        
        Instructions:
        - Say PASS if the AI grading correctly follows the rubric criteria
        - Say FAIL if the AI grading does NOT match the rubric requirements
        - Your decision should be based on whether the AI properly applied the rubric, NOT whether the student deserved a different score
        """
        grading_evaluation = model.generate_content(prompt_evaluate).text
        
        # Parse the evaluation to extract PASS/FAIL and reasoning
        evaluation_status = "Unknown"
        if "PASS" in grading_evaluation:
            evaluation_status = "PASS"
        elif "FAIL" in grading_evaluation:
            evaluation_status = "FAIL"
        
        print(f"Grading Evaluation: {evaluation_status}")
        print(f"Full Response: {grading_evaluation}")
        
        Progress.objects.create(
            student=student_obj,
            rubric=rubric,
            user=request.user,
            score=percentage,
            evaluate=evaluation_status
        )
        return render(request, 'app/grade_result.html', context)
            
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON: {e}")
        print(f"Raw response: {assignment_score}")
        return redirect('dashboard')

# Create your views here.
def home(request):
    if request.user.is_authenticated:
        return render(request, 'app/home_logged.html')
    else:
        return render(request, 'app/home.html')
def dashboard(request):
    print(request.user)
    try:
        rubrics = Rubric.objects.filter(user = request.user)
        students = Student.objects.filter(user = request.user)
    except Exception as e:
        rubrics = []
        students = []    
        print(e)
    
    context = {
        'rubrics': rubrics,
        'students': students
    }
    return render(request, 'app/dashboard.html', context)
@login_required
@require_http_methods(["POST"]) 
def create_rubric(request):
    title = request.POST.get('title', '').strip()
    description = request.POST.get('description', '').strip()
    strictness = request.POST.get('strictness', '').strip()
    subject = request.POST.get('subject', '').strip()
    grade = request.POST.get('grade-level', '').strip()
    print('grade level:', grade)
    if not title:
        rubrics = Rubric.objects.all()
        context = {
            'rubrics': rubrics,
            'error': 'Title is required.'
        }
        return render(request, 'app/dashboard.html', context, status=400)

    Rubric.objects.create(subject = subject, title=title, description=description, strictness=strictness , grade_level=grade, user = request.user)
    return redirect('dashboard')

def progress(request):
    rubrics = Rubric.objects.filter(user=request.user)
    students = Student.objects.filter(user=request.user)
    
    subjects_data = {}
    subject_choices = [
        ('Math', 'Math'),
        ('Science', 'Science'),
        ('English', 'English'),
        ('History', 'History'),
        ('Geography', 'Geography'),
        ('Art', 'Art'),
        ('Music', 'Music'),
    ]
    
    for subject_choice in subject_choices:
        subject = subject_choice[0]
        subject_rubrics = rubrics.filter(subject=subject)
        
        if subject_rubrics.exists():
            # Get all progress data for this subject
            progress_data = Progress.objects.filter(
                rubric__subject=subject,
                user=request.user
            ).order_by('created_at')
            
            # Get all unique students who have progress in this subject
            subject_students = Student.objects.filter(
                progress__rubric__subject=subject,
                progress__user=request.user
            ).distinct()
            
            # Build individual student data with their own labels
            individual_data = {}
            for student in subject_students:
                # Get this student's specific progress
                student_progress = progress_data.filter(student=student).order_by('created_at')
                
                # Build labels and scores for this specific student
                # Group by rubric title only to avoid duplicates
                student_rubrics = student_progress.values_list('rubric__title', flat=True).distinct()
                student_labels = [f"{rubric_title}" for rubric_title in student_rubrics]
                
                # Get average score for each rubric (in case student took same rubric multiple times)
                student_scores = []
                for rubric_title in student_rubrics:
                    avg_score = student_progress.filter(
                        rubric__title=rubric_title
                    ).aggregate(avg_score=Avg('score'))['avg_score']
                    student_scores.append(round(avg_score, 1) if avg_score else None)
                
                individual_data[student.name] = {
                    'id': student.id,
                    'labels': student_labels,
                    'data': student_scores
                }
            
            # For class average, group by rubric title only (not by timestamp)
            # Get unique rubric titles that have been used
            unique_rubrics = progress_data.values_list('rubric__title', flat=True).distinct()
            labels = [f"{rubric_title}" for rubric_title in unique_rubrics]
            
            # Class average data - average all scores for each rubric across all timestamps
            class_scores = []
            for rubric_title in unique_rubrics:
                avg_score = progress_data.filter(
                    rubric__title=rubric_title
                ).aggregate(avg_score=Avg('score'))['avg_score']
                class_scores.append(round(avg_score, 1) if avg_score else None)

            subjects_data[subject] = {
                'labels': labels,  # Overall labels for class average
                'individual': individual_data,  # Each student has their own labels
                'class_average': {
                    'label': 'Class Average',
                    'data': class_scores,
                    'borderColor': 'rgb(75, 192, 192)',
                    'backgroundColor': 'rgba(75, 192, 192, 0.1)',
                    'tension': 0.4,
                    'fill': False
                }
            }
    
    students_data = [{'id': s.id, 'name': s.name} for s in students]
    # Build a list of subjects that actually have data/rubrics for this user
    available_subjects = list(subjects_data.keys())
    context = {
        'subjects_data': json.dumps(subjects_data),
        'students_json': json.dumps(students_data),
        'rubrics': rubrics,
        'students': students,
        'available_subjects': available_subjects
    }
    return render(request, 'app/progress.html', context)


def add_student(request):
    name = request.POST.get('name', '').strip()
    print(name)
    print(request.POST)
    Student.objects.create(name=name ,user = request.user)
    print(Student.objects.all())
    return redirect('dashboard')

def delete_rubric(request, rubric_id):
    try:
        rubric = Rubric.objects.get(id=rubric_id)
        rubric.delete()
        return JsonResponse({'status': 'success'})
    except Rubric.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Rubric not found'}, status=404)
    
def edit_rubric(request, rubric_id):
    rubric = Rubric.objects.get(id=rubric_id)
    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        description = request.POST.get('description', '').strip()
        strictness = request.POST.get('strictness', '').strip()
        grade = request.POST.get('grade-level', '').strip()
        
        rubric.title = title
        rubric.description = description
        rubric.strictness = strictness
        rubric.grade_level = grade
        rubric.save()
        
        return redirect('dashboard')
    
    context = {
        'rubric': rubric
    }
    return redirect('dashboard')

def train_model(request):
    if request.method == 'POST':
        # Process the uploaded files and scores
        files = request.FILES.getlist('files')
        rubric_id = request.POST.get('rubric_for_training')
        grades = request.POST.get('grades', '').strip()
        
        # Get the rubric object
        try:
            rubric = Rubric.objects.get(id=rubric_id, user=request.user)
        except Rubric.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Rubric not found'}, status=404)
        
        # Parse grades from comma-separated string
        grade_list = [int(g.strip()) for g in grades.split(',') if g.strip()]
        
        # Validate that number of files matches number of grades
        if len(files) != len(grade_list):
            return JsonResponse({
                'status': 'error',
                'message': f'Number of files ({len(files)}) must match number of grades ({len(grade_list)})'
            }, status=400)
        
        # Process each file and create training data
        training_data = []
        for i, file in enumerate(files):
            # Read file content based on type
            if file.name.endswith('.docx'):
                doc = Document(io.BytesIO(file.read()))
                file_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            elif file.name.endswith('.pdf'):
                pdf = PdfReader(io.BytesIO(file.read()))
                file_content = '\n'.join([page.extract_text() for page in pdf.pages])
            elif file.name.endswith('.txt'):
                file_content = file.read().decode('utf-8')
            else:
                continue
            
            training_data.append({
                'assignment': file_content,
                'score': grade_list[i]
            })
            
            print(f"Processed file {i+1}: {file.name} - Score: {grade_list[i]}")
        
        print(f"Training model with {len(files)} files for rubric: {rubric.title} (ID: {rubric_id})")
        print(f"Grades received: {grade_list}")
        
        # Store training examples in the rubric's training_data field
        existing_data = rubric.training_data if rubric.training_data else []
        existing_data.extend(training_data)
        rubric.training_data = existing_data
        rubric.save()
        
        messages.success(request, f'Successfully added {len(training_data)} training examples for rubric "{rubric.title}"')
        
        # Redirect back to dashboard with success message
        return redirect('dashboard')
    
    # GET request - show the form
    rubrics = Rubric.objects.filter(user=request.user)
    context = {
        'rubrics': rubrics
    }
    return render(request, 'app/train_model.html', context)